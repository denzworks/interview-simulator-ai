import streamlit as st
import requests
import json
import re
import os
import io
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from datetime import datetime
from typing import Optional

st.set_page_config(
    page_title="AI Mülakat Simülatörü",
    page_icon="⭐",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>.main-header{background:rgba(255,255,255,.05);backdrop-filter:blur(16px);border:1px solid rgba(255,255,255,.08);border-radius:18px;padding:18px 24px;color:white;text-align:center;margin-bottom:20px;box-shadow:0 8px 24px rgba(0,0,0,.18)}.metric-card{background:rgba(255,255,255,.05);backdrop-filter:blur(16px);border:1px solid rgba(255,255,255,.08);border-left:4px solid #6366f1;border-radius:18px;padding:20px;color:white;transition:.25s ease}.metric-card:hover{transform:translateY(-2px);box-shadow:0 10px 30px rgba(99,102,241,.15)}.score-green{color:#22c55e;font-size:1.6rem;font-weight:700}.score-yellow{color:#f59e0b;font-size:1.6rem;font-weight:700}.score-red{color:#ef4444;font-size:1.6rem;font-weight:700}.question-card{background:rgba(255,255,255,.05);backdrop-filter:blur(16px);border:1px solid rgba(255,255,255,.08);border-left:5px solid #6366f1;border-radius:18px;padding:20px;margin-bottom:15px;color:white;box-shadow:0 8px 24px rgba(0,0,0,.15)}.feedback-card{background:rgba(34,197,94,.08);border:1px solid rgba(34,197,94,.15);border-left:4px solid #22c55e;border-radius:18px;padding:18px;margin-top:15px;color:#bbf7d0}.hire-strong,.hire-hire,.hire-borderline,.hire-no{padding:18px;border-radius:18px;font-weight:700;font-size:1.15rem;text-align:center;border:1px solid transparent;backdrop-filter:blur(12px)}.hire-strong{background:rgba(34,197,94,.12);color:#86efac;border-color:rgba(34,197,94,.25)}.hire-hire{background:rgba(59,130,246,.12);color:#93c5fd;border-color:rgba(59,130,246,.25)}.hire-borderline{background:rgba(245,158,11,.12);color:#fde68a;border-color:rgba(245,158,11,.25)}.hire-no{background:rgba(239,68,68,.12);color:#fca5a5;border-color:rgba(239,68,68,.25)}.stTextArea textarea,.stTextInput input{background:rgba(255,255,255,.04)!important;border:1px solid rgba(255,255,255,.08)!important;border-radius:12px!important;color:white!important}.stTextArea textarea:focus,.stTextInput input:focus{border-color:#6366f1!important;box-shadow:0 0 0 3px rgba(99,102,241,.18)!important}.stSelectbox div[data-baseweb="select"]{background:rgba(255,255,255,.04);border-radius:12px;border:1px solid rgba(255,255,255,.08)}section[data-testid="stSidebar"]{background:linear-gradient(180deg,rgba(17,24,39,.95),rgba(15,23,42,.95))}[data-testid="metric-container"]{background:rgba(255,255,255,.04);border:1px solid rgba(255,255,255,.08);border-radius:16px;padding:15px}[data-testid=stToolbar]{display:none!important}#MainMenu{display:none!important}header[data-testid=stHeader]{display:none!important}footer{display:none!important}</style>
""", unsafe_allow_html=True)


def init_session_state() -> None:
    defaults = {
        "page": "home",
        "api_key": "",
        "job_description": "",
        "job_analysis": None,
        "questions": None,
        "all_questions": [],
        "current_question_index": 0,
        "answers": [],
        "scores": [],
        "feedbacks": [],
        "interview_started": False,
        "interview_completed": False,
        "final_scores": None,
        "final_report": None,
        "hiring_recommendation": None,
        "model": "openrouter/auto:free"
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def call_openrouter(prompt: str, system_prompt: str = "") -> str:
    api_key = st.session_state.get("api_key", "") or os.environ.get("OPENROUTER_API_KEY", "sk-or-v1-???") # APİ KEY
    if not api_key:
        st.error("OpenRouter API anahtarı eksik. Lütfen kenar çubuğundan girin.")
        return ""

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://interview-simulator-ai.app",
        "X-Title": "Interview Simulator AI"
    }

    messages = []
    if system_prompt:
        messages.append({"role": "system", "content": system_prompt})
    messages.append({"role": "user", "content": prompt})

    payload = {
        "model": st.session_state.get("model", "openrouter/auto:free"),
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }

    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )
        response.raise_for_status()
        data = response.json()
        return data["choices"][0]["message"]["content"]
    except requests.exceptions.Timeout:
        st.error("İstek zaman aşımına uğradı. Lütfen tekrar deneyin.")
        return ""
    except requests.exceptions.HTTPError as e:
        st.error(f"HTTP Hatası: {e.response.status_code} - {e.response.text}")
        return ""
    except Exception as e:
        st.error(f"OpenRouter bağlantı hatası: {str(e)}")
        return ""


def extract_json_from_text(text: str) -> Optional[dict]:
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    patterns = [
        r'```json\s*([\s\S]*?)\s*```',
        r'```\s*([\s\S]*?)\s*```',
        r'(\{[\s\S]*\})',
        r'(\[[\s\S]*\])'
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL)
        if match:
            try:
                return json.loads(match.group(1))
            except json.JSONDecodeError:
                continue
    return None


def analyze_job_description(job_text: str) -> Optional[dict]:
    system_prompt = (
        "You are an expert HR analyst and recruiter. Analyze job descriptions precisely. "
        "Always respond with valid JSON only, no additional text."
    )
    prompt = f"""Analyze this job description and extract structured information.

Job Description:
{job_text}

Return ONLY valid JSON in this exact format:
{{
    "job_title": "extracted job title",
    "level": "Junior/Mid/Senior/Lead/Principal/Director",
    "required_skills": ["skill1", "skill2", "skill3"],
    "preferred_skills": ["skill1", "skill2"],
    "experience_years": "X-Y years or X+ years",
    "department": "Engineering/Product/Design/etc",
    "key_responsibilities": ["responsibility1", "responsibility2"]
}}"""

    response = call_openrouter(prompt, system_prompt)
    if not response:
        return None

    result = extract_json_from_text(response)
    if not result:
        st.error("İş ilanı analizi ayrıştırılamadı. Lütfen tekrar deneyin.")
        return None
    return result


def generate_interview_questions(job_analysis: dict) -> Optional[dict]:
    system_prompt = (
        "You are an expert technical interviewer with 15+ years of experience. "
        "Generate challenging, relevant interview questions. "
        "Always respond with valid JSON only."
    )

    skills_str = ", ".join(job_analysis.get("required_skills", []))
    level = job_analysis.get("level", "Mid")
    title = job_analysis.get("job_title", "Software Engineer")

    prompt = f"""Generate interview questions for a {level} {title} position.
Required skills: {skills_str}

Return ONLY valid JSON in this exact format:
{{
    "technical_questions": [
        "technical question 1",
        "technical question 2",
        "technical question 3",
        "technical question 4",
        "technical question 5",
        "technical question 6",
        "technical question 7",
        "technical question 8",
        "technical question 9",
        "technical question 10"
    ],
    "behavioral_questions": [
        "behavioral question 1",
        "behavioral question 2",
        "behavioral question 3",
        "behavioral question 4",
        "behavioral question 5"
    ],
    "problem_solving_questions": [
        "problem solving question 1",
        "problem solving question 2",
        "problem solving question 3",
        "problem solving question 4",
        "problem solving question 5"
    ]
}}

Make questions specific to {title} role and {level} level. Questions should be challenging and realistic."""

    response = call_openrouter(prompt, system_prompt)
    if not response:
        return None

    result = extract_json_from_text(response)
    if not result:
        st.error("Mülakat soruları ayrıştırılamadı. Lütfen tekrar deneyin.")
        return None
    return result


def evaluate_answer(
    question: str,
    answer: str,
    category: str,
    job_analysis: dict
) -> Optional[dict]:
    system_prompt = (
        "You are an expert interviewer evaluating candidate responses. "
        "Be objective, fair, and constructive. "
        "Always respond with valid JSON only."
    )

    title = job_analysis.get("job_title", "")
    level = job_analysis.get("level", "")

    prompt = f"""Evaluate this interview answer for a {level} {title} position.

Category: {category}
Question: {question}
Candidate Answer: {answer}

Score each dimension from 1-10 and provide specific feedback.

Return ONLY valid JSON in this exact format:
{{
    "technical_accuracy": 7,
    "clarity": 8,
    "completeness": 6,
    "communication": 7,
    "overall_score": 7,
    "feedback": "Detailed constructive feedback explaining what was good and what could be improved",
    "strengths_in_answer": ["strength1", "strength2"],
    "improvements_needed": ["improvement1", "improvement2"]
}}

Be strict but fair. Consider the {level} level expectations."""

    response = call_openrouter(prompt, system_prompt)
    if not response:
        return None

    result = extract_json_from_text(response)
    if not result:
        return None
    return result


def calculate_final_scores(scores: list, questions: list) -> dict:
    if not scores:
        return {
            "overall": 0,
            "technical": 0,
            "behavioral": 0,
            "problem_solving": 0,
            "communication": 0,
            "total_questions": 0,
            "answered_questions": 0
        }

    technical_scores = []
    behavioral_scores = []
    problem_solving_scores = []
    communication_scores = []
    overall_scores = []

    for i, score in enumerate(scores):
        if score is None:
            continue
        cat = questions[i].get("category", "technical") if i < len(questions) else "technical"
        overall_val = score.get("overall_score", score.get("technical_accuracy", 5))
        comm_val = score.get("communication", 5)
        overall_scores.append(overall_val)
        communication_scores.append(comm_val)
        if cat == "technical":
            technical_scores.append(score.get("technical_accuracy", overall_val))
        elif cat == "behavioral":
            behavioral_scores.append(overall_val)
        elif cat == "problem_solving":
            problem_solving_scores.append(overall_val)

    def safe_avg(lst):
        return round(float(np.mean(lst)) * 10, 1) if lst else 0.0

    return {
        "overall": safe_avg(overall_scores),
        "technical": safe_avg(technical_scores),
        "behavioral": safe_avg(behavioral_scores),
        "problem_solving": safe_avg(problem_solving_scores),
        "communication": safe_avg(communication_scores),
        "total_questions": len(questions),
        "answered_questions": len([s for s in scores if s is not None])
    }


def generate_hiring_recommendation(final_scores: dict, job_analysis: dict) -> str:
    overall = final_scores.get("overall", 0)
    if overall >= 80:
        return "Strong Hire"
    elif overall >= 65:
        return "Hire"
    elif overall >= 45:
        return "Borderline"
    else:
        return "No Hire"


def generate_interview_report(
    job_analysis: dict,
    questions: list,
    answers: list,
    scores: list,
    final_scores: dict,
    hiring_recommendation: str
) -> Optional[dict]:
    system_prompt = (
        "You are an expert HR analyst creating detailed interview performance reports. "
        "Be comprehensive, specific, and actionable. "
        "Always respond with valid JSON only."
    )

    score_summary = json.dumps(final_scores, indent=2)
    qa_pairs = []
    for i, (q, a, s) in enumerate(zip(questions, answers, scores)):
        if a:
            qa_pairs.append(
                f"Q{i+1} ({q.get('category', 'technical')}): {q.get('question', '')}\n"
                f"Answer: {a[:200]}...\n"
                f"Score: {s.get('overall_score', 'N/A') if s else 'Not answered'}"
            )
    qa_summary = "\n\n".join(qa_pairs[:10])

    prompt = f"""Create a comprehensive interview report for this candidate.

Job: {job_analysis.get('job_title', '')} ({job_analysis.get('level', '')})
Hiring Recommendation: {hiring_recommendation}

Final Scores:
{score_summary}

Interview Q&A Summary:
{qa_summary}

Return ONLY valid JSON in this exact format:
{{
    "executive_summary": "2-3 sentence overview of candidate performance",
    "strengths": ["strength1", "strength2", "strength3", "strength4"],
    "weaknesses": ["weakness1", "weakness2", "weakness3"],
    "missing_skills": ["skill1", "skill2", "skill3"],
    "recommendations": ["recommendation1", "recommendation2", "recommendation3"],
    "development_plan": ["development action 1", "development action 2", "development action 3"],
    "key_observations": ["observation1", "observation2", "observation3"],
    "fit_assessment": "Detailed assessment of cultural and technical fit"
}}"""

    response = call_openrouter(prompt, system_prompt)
    if not response:
        return None

    result = extract_json_from_text(response)
    if not result:
        return None
    return result


def export_docx_report(
    job_analysis: dict,
    final_scores: dict,
    report: dict,
    hiring_recommendation: str,
    questions: list,
    answers: list,
    scores: list
) -> bytes:
    doc = Document()

    title_section = doc.sections[0]
    title_section.page_width = Inches(8.5)
    title_section.page_height = Inches(11)

    title = doc.add_heading("Interview Simulator AI", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(102, 126, 234)

    subtitle = doc.add_paragraph("Performans Değerlendirme Raporu")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph(f"Oluşturulma Tarihi: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph("─" * 60)

    doc.add_heading("Pozisyon Bilgileri", 1)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = "Light Grid Accent 1"
    cells_data = [
        ("Pozisyon Adı", job_analysis.get("job_title", "N/A")),
        ("Seviye", job_analysis.get("level", "N/A")),
        ("Gereken Deneyim", job_analysis.get("experience_years", "N/A"))
    ]
    for i, (label, value) in enumerate(cells_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    rec_colors = {
        "Strong Hire": RGBColor(34, 197, 94),
        "Hire": RGBColor(59, 130, 246),
        "Borderline": RGBColor(234, 179, 8),
        "No Hire": RGBColor(239, 68, 68)
    }
    doc.add_heading("İşe Alım Tavsiyesi", 1)
    rec_para = doc.add_paragraph()
    rec_tr_map_docx = {"Strong Hire": "Kesinlikle İşe Al", "Hire": "İşe Al", "Borderline": "Sınırda", "No Hire": "İşe Alma"}
    rec_run = rec_para.add_run(f"▶ {rec_tr_map_docx.get(hiring_recommendation, hiring_recommendation)}")
    rec_run.font.size = Pt(16)
    rec_run.font.bold = True
    rec_run.font.color.rgb = rec_colors.get(hiring_recommendation, RGBColor(0, 0, 0))

    doc.add_heading("Performans Puanları", 1)
    score_table = doc.add_table(rows=5, cols=2)
    score_table.style = "Light Grid Accent 1"
    score_items = [
        ("Genel Puan", final_scores.get("overall", 0)),
        ("Teknik Puan", final_scores.get("technical", 0)),
        ("İletişim Puanı", final_scores.get("communication", 0)),
        ("Problem Çözme Puanı", final_scores.get("problem_solving", 0)),
        ("Davranışsal Puan", final_scores.get("behavioral", 0))
    ]
    for i, (label, score) in enumerate(score_items):
        row = score_table.rows[i]
        row.cells[0].text = label
        score_run = row.cells[1].paragraphs[0].add_run(f"{score}/100")
        score_run.font.bold = True
        if score >= 80:
            score_run.font.color.rgb = RGBColor(34, 197, 94)
        elif score >= 60:
            score_run.font.color.rgb = RGBColor(234, 179, 8)
        else:
            score_run.font.color.rgb = RGBColor(239, 68, 68)

    doc.add_paragraph()
    doc.add_heading("Yönetici Özeti", 1)
    doc.add_paragraph(report.get("executive_summary", ""))

    for section_title, section_key in [
        ("Güçlü Yönler", "strengths"),
        ("Geliştirilmesi Gereken Alanlar", "weaknesses"),
        ("Eksik Beceriler", "missing_skills"),
        ("Öneriler", "recommendations"),
        ("Gelişim Planı", "development_plan")
    ]:
        items = report.get(section_key, [])
        if items:
            doc.add_heading(section_title, 2)
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(item))

    doc.add_heading("Mülakat Soru-Cevap Detayları", 1)
    for i, (q, a, s) in enumerate(zip(questions, answers, scores)):
        if not a:
            continue
        cat_tr = {"technical": "Teknik", "behavioral": "Davranışsal", "problem_solving": "Problem Çözme"}
        cat_label = cat_tr.get(q.get('category', 'technical'), q.get('category', 'technical').title())
        doc.add_heading(f"Soru {i+1} - {cat_label}", 3)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run("S: " + q.get("question", ""))
        q_run.font.bold = True
        doc.add_paragraph("C: " + str(a))
        if s:
            score_val = s.get("overall_score", 0)
            sp = doc.add_paragraph()
            sr = sp.add_run(f"Puan: {score_val}/10")
            sr.font.bold = True
            if score_val >= 8:
                sr.font.color.rgb = RGBColor(34, 197, 94)
            elif score_val >= 6:
                sr.font.color.rgb = RGBColor(234, 179, 8)
            else:
                sr.font.color.rgb = RGBColor(239, 68, 68)
            doc.add_paragraph(f"Geri Bildirim: {s.get('feedback', '')}")
        doc.add_paragraph("─" * 40)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_pdf_report(
    job_analysis: dict,
    final_scores: dict,
    report: dict,
    hiring_recommendation: str,
    questions: list,
    answers: list,
    scores: list
) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=24,
        textColor=colors.HexColor("#667eea"),
        spaceAfter=6,
        alignment=TA_CENTER
    )
    subtitle_style = ParagraphStyle(
        "Subtitle",
        parent=styles["Normal"],
        fontSize=13,
        textColor=colors.grey,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    heading1_style = ParagraphStyle(
        "H1Custom",
        parent=styles["Heading1"],
        fontSize=14,
        textColor=colors.HexColor("#1e293b"),
        spaceBefore=16,
        spaceAfter=8
    )
    body_style = ParagraphStyle(
        "BodyCustom",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#374151"),
        spaceAfter=6,
        leading=14
    )
    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=styles["Normal"],
        fontSize=10,
        leftIndent=20,
        spaceAfter=4,
        bulletIndent=10,
        textColor=colors.HexColor("#374151")
    )

    story = []

    story.append(Paragraph("AI Mülakat Simülatörü", title_style))
    story.append(Paragraph("Performans Değerlendirme Raporu", subtitle_style))
    story.append(Paragraph(f"Oluşturulma Tarihi: {datetime.now().strftime('%d %B %Y, %H:%M')}", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#667eea")))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Pozisyon Bilgileri", heading1_style))
    position_data = [
        ["Alan", "Değer"],
        ["Pozisyon Adı", job_analysis.get("job_title", "N/A")],
        ["Seviye", job_analysis.get("level", "N/A")],
        ["Deneyim", job_analysis.get("experience_years", "N/A")]
    ]
    pos_table = Table(position_data, colWidths=[2.5 * inch, 4 * inch])
    pos_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#667eea")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#f8fafc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 10)
    ]))
    story.append(pos_table)
    story.append(Spacer(1, 16))

    rec_color_map = {
        "Strong Hire": colors.HexColor("#22c55e"),
        "Hire": colors.HexColor("#3b82f6"),
        "Borderline": colors.HexColor("#eab308"),
        "No Hire": colors.HexColor("#ef4444")
    }
    rec_bg_map = {
        "Strong Hire": colors.HexColor("#dcfce7"),
        "Hire": colors.HexColor("#dbeafe"),
        "Borderline": colors.HexColor("#fef9c3"),
        "No Hire": colors.HexColor("#fee2e2")
    }
    story.append(Paragraph("İşe Alım Tavsiyesi", heading1_style))
    rec_tr_map_pdf = {"Strong Hire": "Kesinlikle İşe Al", "Hire": "İşe Al", "Borderline": "Sınırda", "No Hire": "İşe Alma"}
    rec_color = rec_color_map.get(hiring_recommendation, colors.black)
    rec_bg = rec_bg_map.get(hiring_recommendation, colors.white)
    rec_style = ParagraphStyle(
        "RecStyle",
        parent=styles["Normal"],
        fontSize=16,
        textColor=rec_color,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold"
    )
    rec_table = Table([[Paragraph(f"▶  {rec_tr_map_pdf.get(hiring_recommendation, hiring_recommendation)}", rec_style)]], colWidths=[6.5 * inch])
    rec_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), rec_bg),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
        ("ROUNDEDCORNERS", (0, 0), (-1, -1), 8)
    ]))
    story.append(rec_table)
    story.append(Spacer(1, 16))

    story.append(Paragraph("Performans Puanları", heading1_style))
    score_data = [["Kategori", "Puan", "Durum"]]
    score_items_pdf = [
        ("Genel Puan", final_scores.get("overall", 0)),
        ("Teknik Puan", final_scores.get("technical", 0)),
        ("İletişim Puanı", final_scores.get("communication", 0)),
        ("Problem Çözme Puanı", final_scores.get("problem_solving", 0)),
        ("Davranışsal Puan", final_scores.get("behavioral", 0))
    ]
    for label, score in score_items_pdf:
        if score >= 80:
            status = "Mükemmel"
        elif score >= 60:
            status = "İyi"
        elif score >= 40:
            status = "Geliştirilmeli"
        else:
            status = "Yetersiz"
        score_data.append([label, f"{score}/100", status])

    s_table = Table(score_data, colWidths=[3 * inch, 1.5 * inch, 2 * inch])
    s_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("FONTNAME", (1, 1), (1, -1), "Helvetica-Bold")
    ]))
    story.append(s_table)
    story.append(Spacer(1, 16))

    story.append(Paragraph("Yönetici Özeti", heading1_style))
    story.append(Paragraph(report.get("executive_summary", ""), body_style))
    story.append(Spacer(1, 12))

    for section_title, section_key in [
        ("Güçlü Yönler", "strengths"),
        ("Geliştirilmesi Gereken Alanlar", "weaknesses"),
        ("Eksik Beceriler", "missing_skills"),
        ("Öneriler", "recommendations"),
        ("Gelişim Planı", "development_plan")
    ]:
        items = report.get(section_key, [])
        if items:
            story.append(Paragraph(section_title, heading1_style))
            for item in items:
                story.append(Paragraph(f"• {item}", bullet_style))
            story.append(Spacer(1, 8))

    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#667eea")))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Mülakat Soru-Cevap Detayları", heading1_style))

    for i, (q, a, s) in enumerate(zip(questions, answers, scores)):
        if not a:
            continue
        cat_tr_pdf = {"technical": "Teknik", "behavioral": "Davranışsal", "problem_solving": "Problem Çözme"}
        cat_label_pdf = cat_tr_pdf.get(q.get("category", "technical"), q.get("category", "technical").title())
        q_style = ParagraphStyle(
            "QStyle",
            parent=styles["Normal"],
            fontSize=10,
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#1e293b"),
            spaceAfter=4
        )
        story.append(Paragraph(f"Soru {i+1} ({cat_label_pdf})", q_style))
        story.append(Paragraph(f"<b>S:</b> {q.get('question', '')}", body_style))
        story.append(Paragraph(f"<b>C:</b> {str(a)[:500]}", body_style))
        if s:
            score_val = s.get("overall_score", 0)
            if score_val >= 8:
                sc_color = "#22c55e"
            elif score_val >= 6:
                sc_color = "#eab308"
            else:
                sc_color = "#ef4444"
            story.append(Paragraph(
                f'<font color="{sc_color}"><b>Puan: {score_val}/10</b></font>',
                body_style
            ))
            story.append(Paragraph(f"<i>Geri Bildirim: {s.get('feedback', '')}</i>", body_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0")))
        story.append(Spacer(1, 8))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def get_score_color(score: float) -> str:
    if score >= 80:
        return "score-green"
    elif score >= 60:
        return "score-yellow"
    else:
        return "score-red"


def render_score_metric(label: str, value: float) -> None:
    css_class = get_score_color(value)
    if value >= 80:
        delta_color = "normal"
    elif value >= 60:
        delta_color = "off"
    else:
        delta_color = "inverse"
    st.metric(label=label, value=f"{value}/100")


def create_radar_chart(final_scores: dict) -> go.Figure:
    categories = ["Teknik", "İletişim", "Problem Çözme", "Davranışsal", "Genel"]
    values = [
        final_scores.get("technical", 0),
        final_scores.get("communication", 0),
        final_scores.get("problem_solving", 0),
        final_scores.get("behavioral", 0),
        final_scores.get("overall", 0)
    ]
    values_normalized = [v / 10 for v in values]
    values_normalized.append(values_normalized[0])
    cats = categories + [categories[0]]

    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=values_normalized,
        theta=cats,
        fill="toself",
        fillcolor="rgba(102, 126, 234, 0.2)",
        line=dict(color="#667eea", width=2),
        name="Performance"
    ))
    fig.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 10], tickfont=dict(size=10))
        ),
        showlegend=False,
        title=dict(text="Performans Radar Grafiği", font=dict(size=16, color="#1e293b")),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        height=350
    )
    return fig


def create_score_distribution_chart(scores: list, questions: list) -> go.Figure:
    q_nums = []
    overall_vals = []
    categories_list = []

    for i, (s, q) in enumerate(zip(scores, questions)):
        if s is None:
            continue
        q_nums.append(f"Q{i+1}")
        overall_vals.append(s.get("overall_score", 0))
        categories_list.append(q.get("category", "technical").title())

    if not q_nums:
        fig = go.Figure()
        fig.add_annotation(text="Henüz veri yok", xref="paper", yref="paper", x=0.5, y=0.5, showarrow=False)
        return fig

    color_map = {"Technical": "#667eea", "Behavioral": "#22c55e", "Problem_Solving": "#f59e0b"}
    bar_colors = [color_map.get(c, "#667eea") for c in categories_list]

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=q_nums,
        y=overall_vals,
        marker_color=bar_colors,
        text=[f"{v}/10" for v in overall_vals],
        textposition="outside",
        hovertemplate="<b>%{x}</b><br>Score: %{y}/10<extra></extra>"
    ))
    fig.add_hline(y=7, line_dash="dash", line_color="#22c55e", annotation_text="İyi (7)")
    fig.add_hline(y=5, line_dash="dash", line_color="#eab308", annotation_text="Ortalama (5)")
    fig.update_layout(
        title=dict(text="Soru Bazlı Puan Dağılımı", font=dict(size=16, color="#1e293b")),
        xaxis_title="Sorular",
        yaxis_title="Puan (1-10)",
        yaxis=dict(range=[0, 11]),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        height=350,
        showlegend=False
    )
    return fig


def create_progress_chart(current_index: int, total: int) -> go.Figure:
    answered = current_index
    remaining = total - current_index
    fig = go.Figure(data=[go.Pie(
        values=[answered, remaining],
        labels=["Tamamlandı", "Kalan"],
        hole=0.65,
        marker_colors=["#667eea", "#e2e8f0"],
        textinfo="none",
        hovertemplate="%{label}: %{value}<extra></extra>"
    )])
    fig.add_annotation(
        text=f"{answered}/{total}",
        x=0.5, y=0.5,
        font_size=20,
        font_color="#1e293b",
        showarrow=False
    )
    fig.update_layout(
        title=dict(text="Mülakat İlerlemesi", font=dict(size=14, color="#1e293b")),
        showlegend=True,
        paper_bgcolor="rgba(0,0,0,0)",
        height=280,
        margin=dict(t=40, b=20, l=20, r=20)
    )
    return fig


def render_home_page() -> None:
    st.markdown("""
    <div class="main-header">
        <h1 style="margin:0; font-size:2.5rem;">⭐ AI Mülakat Simülatörü</h1>
        <p style="margin:0.5rem 0 0 0; font-size:1.1rem; opacity:0.9;">
            Yapay zeka destekli gerçekçi mülakatlar, anlık değerlendirme ve detaylı raporlar
        </p>
    </div>
    """, unsafe_allow_html=True)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.markdown("""
        <div class="metric-card">
            <div style="font-size:2rem;">📄</div>
            <div style="font-weight:600; margin-top:0.5rem;">İlan Analizi</div>
            <div style="color:#64748b; font-size:0.85rem;">AI destekli iş ilanı ayrıştırma</div>
        </div>
        """, unsafe_allow_html=True)
    with col2:
        st.markdown("""
        <div class="metric-card">
            <div style="font-size:2rem;">🤖</div>
            <div style="font-weight:600; margin-top:0.5rem;">20 Soru</div>
            <div style="color:#64748b; font-size:0.85rem;">Teknik, Davranışsal & Problem Çözme</div>
        </div>
        """, unsafe_allow_html=True)
    with col3:
        st.markdown("""
        <div class="metric-card">
            <div style="font-size:2rem;">📊</div>
            <div style="font-weight:600; margin-top:0.5rem;">Anlık Puanlama</div>
            <div style="color:#64748b; font-size:0.85rem;">Gerçek zamanlı cevap değerlendirme</div>
        </div>
        """, unsafe_allow_html=True)
    with col4:
        st.markdown("""
        <div class="metric-card">
            <div style="font-size:2rem;">📋</div>
            <div style="font-weight:600; margin-top:0.5rem;">Tam Rapor</div>
            <div style="color:#64748b; font-size:0.85rem;">PDF & DOCX dışa aktarma</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    col_left, col_right = st.columns([3, 2])

    with col_left:
        st.subheader("📝 İş İlanını Girin")
        st.text_area(
            "İş ilanını buraya yapıştırın:",
            height=300,
            placeholder="Sorumluluklar, gereksinimler ve nitelikleri içeren tam iş ilanını yapıştırın...",
            key="job_desc_input"
        )

        st.subheader("📎 Veya Dosya Yükleyin")
        uploaded_file = st.file_uploader(
            "PDF veya DOCX iş ilanı yükleyin",
            type=["pdf", "docx", "txt"],
            key="file_uploader"
        )

        if uploaded_file is not None:
            if uploaded_file.type == "text/plain":
                text_content = uploaded_file.read().decode("utf-8", errors="ignore")
                st.session_state["job_description"] = text_content
                st.success(f"✅ TXT dosyasından {len(text_content)} karakter yüklendi.")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(uploaded_file)
                text_content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                st.session_state["job_description"] = text_content
                st.success(f"✅ DOCX dosyasından {len(text_content)} karakter yüklendi.")
            elif uploaded_file.type == "application/pdf":
                try:
                    import pdfplumber
                    with pdfplumber.open(uploaded_file) as pdf:
                        text_content = "\n".join([page.extract_text() or "" for page in pdf.pages])
                    st.session_state["job_description"] = text_content
                    st.success(f"✅ PDF dosyasından {len(text_content)} karakter yüklendi.")
                except ImportError:
                    raw_bytes = uploaded_file.read()
                    text_content = raw_bytes.decode("latin-1", errors="ignore")
                    pdf_text = re.sub(r'[^\x20-\x7E\n]', ' ', text_content)
                    pdf_text = re.sub(r'\s+', ' ', pdf_text)
                    st.session_state["job_description"] = pdf_text
                    st.warning("PDF basit ayrıştırıcı ile okundu. Daha iyi sonuç için pdfplumber kurun.")

    with col_right:
        st.subheader("⚙️ Mülakat Ayarları")

        if st.session_state.get("job_analysis"):
            analysis = st.session_state["job_analysis"]
            st.success("✅ İş ilanı analiz edildi!")
            st.markdown(f"""
            <div class="metric-card">
                <div><b>Pozisyon:</b> {analysis.get('job_title', 'N/A')}</div>
                <div><b>Seviye:</b> {analysis.get('level', 'N/A')}</div>
                <div><b>Deneyim:</b> {analysis.get('experience_years', 'N/A')}</div>
            </div>
            """, unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)
            req_skills = analysis.get("required_skills", [])
            if req_skills:
                st.markdown("**Zorunlu Beceriler:**")
                skills_html = " ".join([
                    f'<span style="background:#667eea20; color:#667eea; padding:2px 8px; border-radius:12px; font-size:0.8rem; margin:2px; display:inline-block;">{s}</span>'
                    for s in req_skills[:10]
                ])
                st.markdown(skills_html, unsafe_allow_html=True)

        jd_current = st.session_state.get("job_desc_input", "") or st.session_state.get("job_description", "")
        can_start = bool(jd_current and jd_current.strip() and len(jd_current.strip()) > 50)

        if can_start and not st.session_state.get("job_analysis"):
            if st.button("🔍 İş İlanını Analiz Et", use_container_width=True, type="secondary"):
                with st.spinner("İş ilanı analiz ediliyor..."):
                    analysis = analyze_job_description(jd_current)
                    if analysis:
                        st.session_state["job_analysis"] = analysis
                        st.rerun()

        if st.session_state.get("job_analysis") and st.session_state.get("api_key"):
            st.markdown("---")
            if st.button("🚀 Mülakatı Başlat", use_container_width=True, type="primary"):
                with st.spinner("Mülakat soruları oluşturuluyor..."):
                    questions_data = generate_interview_questions(st.session_state["job_analysis"])
                    if questions_data:
                        all_questions = []
                        for q in questions_data.get("technical_questions", []):
                            all_questions.append({"question": q, "category": "technical"})
                        for q in questions_data.get("behavioral_questions", []):
                            all_questions.append({"question": q, "category": "behavioral"})
                        for q in questions_data.get("problem_solving_questions", []):
                            all_questions.append({"question": q, "category": "problem_solving"})

                        st.session_state["questions"] = questions_data
                        st.session_state["all_questions"] = all_questions
                        st.session_state["current_question_index"] = 0
                        st.session_state["answers"] = [""] * len(all_questions)
                        st.session_state["scores"] = [None] * len(all_questions)
                        st.session_state["feedbacks"] = [None] * len(all_questions)
                        st.session_state["interview_started"] = True
                        st.session_state["interview_completed"] = False
                        st.session_state["page"] = "interview"
                        st.rerun()
        elif not st.session_state.get("api_key"):
            st.warning("⚠️ Mülakatı başlatmak için kenar çubuğundan OpenRouter API anahtarınızı girin.")
        elif not st.session_state.get("job_analysis"):
            st.info("ℹ️ Önce iş ilanını analiz edin.")


def render_interview_page() -> None:
    all_questions = st.session_state.get("all_questions", [])
    current_idx = st.session_state.get("current_question_index", 0)
    answers = st.session_state.get("answers", [])
    scores = st.session_state.get("scores", [])
    job_analysis = st.session_state.get("job_analysis", {})

    if not all_questions:
        st.error("Soru bulunamadı. Lütfen mülakatı yeniden başlatın.")
        if st.button("← Ana Sayfaya Dön"):
            st.session_state["page"] = "home"
            st.rerun()
        return

    total = len(all_questions)

    st.markdown("""
    <div style="background:linear-gradient(135deg,#667eea,#764ba2); padding:1.5rem; border-radius:12px; color:white; margin-bottom:1.5rem;">
        <h2 style="margin:0;">🎤 Mülakat Devam Ediyor</h2>
        <p style="margin:0.3rem 0 0 0; opacity:0.9;">Her soruyu dikkatlice yanıtlayın. Acele etmeyin.</p>
    </div>
    """, unsafe_allow_html=True)

    prog_col, chart_col = st.columns([2, 1])
    with prog_col:
        progress_val = current_idx / total if total > 0 else 0
        st.progress(progress_val)
        st.caption(f"Soru {min(current_idx + 1, total)} / {total}")

        answered_count = len([a for a in answers if a and a.strip()])
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("Toplam Soru", total)
        with c2:
            st.metric("Cevaplanan", answered_count)
        with c3:
            scored_list = [s for s in scores if s is not None]
            if scored_list:
                avg = np.mean([s.get("overall_score", 0) for s in scored_list])
                st.metric("Ort. Puan", f"{avg:.1f}/10")
            else:
                st.metric("Ort. Puan", "—")

    with chart_col:
        fig_prog = create_progress_chart(answered_count, total)
        st.plotly_chart(fig_prog, use_container_width=True, key="progress_chart")

    if current_idx < total:
        current_q = all_questions[current_idx]
        cat = current_q.get("category", "technical")
        cat_labels = {
            "technical": ("⚙️ Teknik", "#667eea"),
            "behavioral": ("🤝 Davranışsal", "#22c55e"),
            "problem_solving": ("🧠 Problem Çözme", "#f59e0b")
        }
        cat_label, cat_color = cat_labels.get(cat, ("❓ Genel", "#64748b"))

        st.markdown(f"""
        <div class="question-card">
            <div style="display:flex; align-items:center; gap:0.5rem; margin-bottom:0.8rem;">
                <span style="background:{cat_color}20; color:{cat_color}; padding:3px 12px; border-radius:20px; font-size:0.85rem; font-weight:600;">{cat_label}</span>
                <span style="color:#cbd5e1; font-size:0.85rem;">Soru {current_idx + 1} / {total}</span>
            </div>
            <div style="font-size:1.1rem; font-weight:500; color:#ffffff; line-height:1.6;">
                {current_q.get("question", "")}
            </div>
        </div>
        """, unsafe_allow_html=True)

        answer_key = f"answer_input_{current_idx}"
        user_answer = st.text_area(
            "Cevabınız:",
            value=answers[current_idx] if current_idx < len(answers) else "",
            height=200,
            placeholder="Detaylı cevabınızı buraya yazın. Mümkün olduğunca somut örnekler verin...",
            key=answer_key
        )

        btn_col1, btn_col2, btn_col3 = st.columns([1, 1, 2])
        with btn_col1:
            if current_idx > 0:
                if st.button("← Önceki", use_container_width=True):
                    if current_idx < len(answers) and user_answer:
                        st.session_state["answers"][current_idx] = user_answer
                    st.session_state["current_question_index"] = current_idx - 1
                    st.rerun()

        with btn_col2:
            skip_label = "Atla →" if current_idx < total - 1 else "Atla (son)"
            if st.button(skip_label, use_container_width=True):
                if current_idx < len(answers):
                    st.session_state["answers"][current_idx] = ""
                if current_idx + 1 < total:
                    st.session_state["current_question_index"] = current_idx + 1
                    st.rerun()

        with btn_col3:
            submit_label = "Gönder & Sonraki →" if current_idx < total - 1 else "Gönder & Bitir 🎉"
            if st.button(submit_label, use_container_width=True, type="primary"):
                if not user_answer or not user_answer.strip():
                    st.warning("Göndermeden önce lütfen bir cevap yazın.")
                else:
                    st.session_state["answers"][current_idx] = user_answer
                    with st.spinner("Cevabınız değerlendiriliyor..."):
                        score_result = evaluate_answer(
                            question=current_q.get("question", ""),
                            answer=user_answer,
                            category=cat,
                            job_analysis=job_analysis
                        )
                        if score_result:
                            st.session_state["scores"][current_idx] = score_result
                            st.session_state["feedbacks"][current_idx] = score_result.get("feedback", "")

                    if current_idx + 1 < total:
                        st.session_state["current_question_index"] = current_idx + 1
                        st.rerun()
                    else:
                        st.session_state["interview_completed"] = True
                        st.session_state["page"] = "analysis"
                        with st.spinner("Nihai puanlar hesaplanıyor..."):
                            final_scores = calculate_final_scores(
                                st.session_state["scores"],
                                st.session_state["all_questions"]
                            )
                            st.session_state["final_scores"] = final_scores
                            rec = generate_hiring_recommendation(final_scores, job_analysis)
                            st.session_state["hiring_recommendation"] = rec
                        st.rerun()

        if scores[current_idx] is not None:
            s = scores[current_idx]
            st.markdown(f"""
            <div class="feedback-card">
                <div style="font-weight:600; color:#166534; margin-bottom:0.5rem;">
                    ✅ Önceki Puan: {s.get("overall_score", "N/A")}/10
                </div>
                <div style="color:#374151; font-size:0.9rem;">{s.get("feedback", "")}</div>
            </div>
            """, unsafe_allow_html=True)

    else:
        st.success("🎉 Tüm soruları tamamladınız! Analiz sayfasına yönlendiriliyorsunuz...")
        st.session_state["interview_completed"] = True
        st.session_state["page"] = "analysis"
        if not st.session_state.get("final_scores"):
            final_scores = calculate_final_scores(
                st.session_state["scores"],
                st.session_state["all_questions"]
            )
            st.session_state["final_scores"] = final_scores
            rec = generate_hiring_recommendation(final_scores, job_analysis)
            st.session_state["hiring_recommendation"] = rec
        st.rerun()


def render_analysis_page() -> None:
    final_scores = st.session_state.get("final_scores")
    job_analysis = st.session_state.get("job_analysis", {})
    hiring_rec = st.session_state.get("hiring_recommendation", "Borderline")
    all_questions = st.session_state.get("all_questions", [])
    answers = st.session_state.get("answers", [])
    scores = st.session_state.get("scores", [])

    if not final_scores:
        st.error("Analiz verisi bulunamadı. Lütfen önce mülakatı tamamlayın.")
        if st.button("← Ana Sayfa"):
            st.session_state["page"] = "home"
            st.rerun()
        return

    st.markdown("""
    <div style="background:linear-gradient(135deg,#1e293b,#334155); padding:1.5rem; border-radius:12px; color:white; margin-bottom:1.5rem;">
        <h2 style="margin:0;">📊 Performans Analizi</h2>
        <p style="margin:0.3rem 0 0 0; opacity:0.9;">Mülakat performansınızın kapsamlı dökümü</p>
    </div>
    """, unsafe_allow_html=True)

    st.subheader("🏆 Puan Özeti")
    m1, m2, m3, m4 = st.columns(4)
    score_map = [
        (m1, "Genel Puan", final_scores.get("overall", 0)),
        (m2, "Teknik Puan", final_scores.get("technical", 0)),
        (m3, "İletişim", final_scores.get("communication", 0)),
        (m4, "Problem Çözme", final_scores.get("problem_solving", 0))
    ]
    for col, label, val in score_map:
        with col:
            if val >= 80:
                delta = "Mükemmel"
            elif val >= 60:
                delta = "İyi"
            else:
                delta = "Geliştirilmeli"
            st.metric(label=label, value=f"{val}/100", delta=delta)

    st.markdown("---")

    chart_col1, chart_col2 = st.columns(2)
    with chart_col1:
        radar = create_radar_chart(final_scores)
        st.plotly_chart(radar, use_container_width=True, key="radar_chart")
    with chart_col2:
        dist = create_score_distribution_chart(scores, all_questions)
        st.plotly_chart(dist, use_container_width=True, key="dist_chart")

    st.markdown("---")

    if not st.session_state.get("final_report"):
        with st.spinner("AI ile detaylı rapor oluşturuluyor..."):
            report = generate_interview_report(
                job_analysis=job_analysis,
                questions=all_questions,
                answers=answers,
                scores=scores,
                final_scores=final_scores,
                hiring_recommendation=hiring_rec
            )
            if report:
                st.session_state["final_report"] = report

    report = st.session_state.get("final_report", {})

    if report:
        rep_col1, rep_col2, rep_col3 = st.columns(3)
        with rep_col1:
            st.subheader("💪 Güçlü Yönler")
            for s in report.get("strengths", []):
                st.markdown(f"✅ {s}")
        with rep_col2:
            st.subheader("⚠️ Zayıf Yönler")
            for w in report.get("weaknesses", []):
                st.markdown(f"⚠️ {w}")
        with rep_col3:
            st.subheader("📚 Eksik Beceriler")
            for m in report.get("missing_skills", []):
                st.markdown(f"❌ {m}")

    st.markdown("---")
    st.subheader("📝 Soru-Cevap İncelemesi")
    for i, (q, a, s) in enumerate(zip(all_questions, answers, scores)):
        if not a:
            continue
        cat = q.get("category", "technical")
        cat_emoji = {"technical": "⚙️", "behavioral": "🤝", "problem_solving": "🧠"}.get(cat, "❓")
        with st.expander(f"{cat_emoji} S{i+1}: {q.get('question', '')[:80]}..."):
            st.markdown(f"**Soru:** {q.get('question', '')}")
            st.markdown(f"**Cevabınız:** {a}")
            if s:
                scol1, scol2, scol3, scol4 = st.columns(4)
                with scol1:
                    st.metric("Teknik Doğruluk", f"{s.get('technical_accuracy', 0)}/10")
                with scol2:
                    st.metric("Açıklık", f"{s.get('clarity', 0)}/10")
                with scol3:
                    st.metric("Eksiksizlik", f"{s.get('completeness', 0)}/10")
                with scol4:
                    st.metric("İletişim", f"{s.get('communication', 0)}/10")
                st.info(f"💬 **Geri Bildirim:** {s.get('feedback', '')}")

    st.markdown("---")
    if st.button("📋 Tam Raporu Görüntüle & İndir", use_container_width=True, type="primary"):
        st.session_state["page"] = "report"
        st.rerun()


def render_report_page() -> None:
    final_scores = st.session_state.get("final_scores", {})
    job_analysis = st.session_state.get("job_analysis", {})
    hiring_rec = st.session_state.get("hiring_recommendation", "Borderline")
    report = st.session_state.get("final_report", {})
    all_questions = st.session_state.get("all_questions", [])
    answers = st.session_state.get("answers", [])
    scores = st.session_state.get("scores", [])

    if not final_scores or not report:
        st.error("Rapor verisi mevcut değil. Lütfen önce mülakatı tamamlayın.")
        if st.button("← Ana Sayfa"):
            st.session_state["page"] = "home"
            st.rerun()
        return

    st.markdown("""
    <div style="background:linear-gradient(135deg,#0f172a,#1e293b); padding:1.5rem; border-radius:12px; color:white; margin-bottom:1.5rem;">
        <h2 style="margin:0;">📋 Mülakat Raporu</h2>
        <p style="margin:0.3rem 0 0 0; opacity:0.9;">İşe alım tavsiyesi ile birlikte eksiksiz performans raporu</p>
    </div>
    """, unsafe_allow_html=True)

    st.subheader("⭐ İşe Alım Tavsiyesi")
    rec_css_map = {
        "Strong Hire": "hire-strong",
        "Hire": "hire-hire",
        "Borderline": "hire-borderline",
        "No Hire": "hire-no"
    }
    rec_icon_map = {
        "Strong Hire": "🟢",
        "Hire": "🔵",
        "Borderline": "🟡",
        "No Hire": "🔴"
    }
    rec_tr_map = {
        "Strong Hire": "Kesinlikle İşe Al",
        "Hire": "İşe Al",
        "Borderline": "Sınırda",
        "No Hire": "İşe Alma"
    }
    css_class = rec_css_map.get(hiring_rec, "hire-borderline")
    icon = rec_icon_map.get(hiring_rec, "⚪")
    rec_tr = rec_tr_map.get(hiring_rec, hiring_rec)
    st.markdown(f'<div class="{css_class}">{icon} {rec_tr}</div>', unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    if report.get("executive_summary"):
        st.subheader("📌 Yönetici Özeti")
        st.info(report.get("executive_summary", ""))

    st.markdown("---")

    r_col1, r_col2 = st.columns(2)
    with r_col1:
        st.subheader("💪 Güçlü Yönler")
        for s in report.get("strengths", []):
            st.markdown(f"✅ {s}")

        st.markdown("<br>", unsafe_allow_html=True)
        st.subheader("📚 Eksik Beceriler")
        for m in report.get("missing_skills", []):
            st.markdown(f"❌ {m}")

    with r_col2:
        st.subheader("⚠️ Geliştirilmesi Gereken Alanlar")
        for w in report.get("weaknesses", []):
            st.markdown(f"⚠️ {w}")

        st.markdown("<br>", unsafe_allow_html=True)
        st.subheader("💡 Öneriler")
        for rec in report.get("recommendations", []):
            st.markdown(f"➡️ {rec}")

    if report.get("development_plan"):
        st.markdown("---")
        st.subheader("🗺️ Gelişim Planı")
        for i, item in enumerate(report.get("development_plan", []), 1):
            st.markdown(f"**{i}.** {item}")

    if report.get("fit_assessment"):
        st.markdown("---")
        st.subheader("🔍 Uyum Değerlendirmesi")
        st.write(report.get("fit_assessment", ""))

    st.markdown("---")
    st.subheader("📊 Nihai Puan Özeti")
    scores_df = pd.DataFrame([
        {"Kategori": "Genel", "Puan": final_scores.get("overall", 0), "Maks": 100},
        {"Kategori": "Teknik", "Puan": final_scores.get("technical", 0), "Maks": 100},
        {"Kategori": "İletişim", "Puan": final_scores.get("communication", 0), "Maks": 100},
        {"Kategori": "Problem Çözme", "Puan": final_scores.get("problem_solving", 0), "Maks": 100},
        {"Kategori": "Davranışsal", "Puan": final_scores.get("behavioral", 0), "Maks": 100}
    ])
    scores_df["Not"] = scores_df["Puan"].apply(
        lambda x: "A" if x >= 90 else ("B" if x >= 80 else ("C" if x >= 70 else ("D" if x >= 60 else "F")))
    )
    st.dataframe(scores_df, use_container_width=True, hide_index=True)

    st.markdown("---")
    st.subheader("⬇️ Raporu İndir")
    dl_col1, dl_col2 = st.columns(2)

    with dl_col1:
        with st.spinner("DOCX oluşturuluyor..."):
            docx_bytes = export_docx_report(
                job_analysis=job_analysis,
                final_scores=final_scores,
                report=report,
                hiring_recommendation=hiring_rec,
                questions=all_questions,
                answers=answers,
                scores=scores
            )
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        title_safe = re.sub(r'[^a-zA-Z0-9_]', '_', job_analysis.get("job_title", "interview"))
        st.download_button(
            label="📄 DOCX Raporu İndir",
            data=docx_bytes,
            file_name=f"mulakat_raporu_{title_safe}_{timestamp}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    with dl_col2:
        with st.spinner("PDF oluşturuluyor..."):
            pdf_bytes = export_pdf_report(
                job_analysis=job_analysis,
                final_scores=final_scores,
                report=report,
                hiring_recommendation=hiring_rec,
                questions=all_questions,
                answers=answers,
                scores=scores
            )
        st.download_button(
            label="📑 PDF Raporu İndir",
            data=pdf_bytes,
            file_name=f"mulakat_raporu_{title_safe}_{timestamp}.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🏠 Yeni Mülakat Başlat", use_container_width=True):
        for key in list(st.session_state.keys()):
            if key not in ["api_key", "model"]:
                del st.session_state[key]
        init_session_state()
        st.session_state["page"] = "home"
        st.rerun()


def render_sidebar() -> None:
    with st.sidebar:
        st.markdown("""
        <div style="background:linear-gradient(135deg,#667eea,#764ba2); padding:1rem; border-radius:10px; text-align:center; margin-bottom:1rem;">
            <div style="font-size:2rem;">⭐</div>
            <div style="color:white; font-weight:700; font-size:1.1rem;">AI Mülakat Simülatörü</div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("### 🔑 API Yapılandırması")
        api_key_input = st.text_input(
            "OpenRouter API Anahtarı",
            value=st.session_state.get("api_key", ""),
            type="password",
            placeholder="sk-or-...",
            key="api_key_sidebar"
        )
        if api_key_input:
            st.session_state["api_key"] = api_key_input

        env_key = os.environ.get("OPENROUTER_API_KEY", "")
        if env_key and not st.session_state.get("api_key"):
            st.session_state["api_key"] = env_key
            st.success("✅ API anahtarı ortam değişkeninden yüklendi.")
        elif st.session_state.get("api_key"):
            st.success("✅ API anahtarı ayarlandı.")
        else:
            st.warning("⚠️ API anahtarı gerekli.")

        model_options = [
            "openrouter/auto:free",
            "mistralai/mistral-7b-instruct:free",
            "google/gemma-2-9b-it:free",
            "meta-llama/llama-3.1-8b-instruct:free",
            "microsoft/phi-3-mini-128k-instruct:free",
            "anthropic/claude-3-haiku",
            "openai/gpt-3.5-turbo",
            "openai/gpt-4o-mini"
        ]
        selected_model = st.selectbox(
            "AI Modeli",
            options=model_options,
            index=0,
            key="model_select"
        )
        st.session_state["model"] = selected_model

        st.markdown("---")
        st.markdown("### 🧭 Gezinme")
        current_page = st.session_state.get("page", "home")

        nav_items = [
            ("🏠 Ana Sayfa", "home", True),
            ("🎤 Mülakat", "interview", st.session_state.get("interview_started", False)),
            ("📊 Analiz", "analysis", st.session_state.get("interview_completed", False)),
            ("📋 Rapor", "report", bool(st.session_state.get("final_report")))
        ]
        for label, page_name, enabled in nav_items:
            if enabled:
                btn_type = "primary" if current_page == page_name else "secondary"
                if st.button(label, use_container_width=True, key=f"nav_{page_name}", type=btn_type):
                    st.session_state["page"] = page_name
                    st.rerun()
            else:
                st.button(label, use_container_width=True, disabled=True, key=f"nav_{page_name}_disabled")

        st.markdown("---")
        if st.session_state.get("job_analysis"):
            analysis = st.session_state["job_analysis"]
            st.markdown("### 📄 İlan Bilgisi")
            st.markdown(f"**Başlık:** {analysis.get('job_title', 'N/A')}")
            st.markdown(f"**Seviye:** {analysis.get('level', 'N/A')}")

        if st.session_state.get("final_scores"):
            fs = st.session_state["final_scores"]
            st.markdown("### 📊 Hızlı Puanlar")
            overall = fs.get("overall", 0)
            if overall >= 80:
                color = "🟢"
            elif overall >= 60:
                color = "🟡"
            else:
                color = "🔴"
            st.markdown(f"{color} **Genel:** {overall}/100")
            st.markdown(f"⚙️ **Teknik:** {fs.get('technical', 0)}/100")
            st.markdown(f"💬 **İletişim:** {fs.get('communication', 0)}/100")

        if st.session_state.get("hiring_recommendation"):
            st.markdown("---")
            st.markdown("### ⭐ İşe Alım Tavsiyesi")
            rec = st.session_state["hiring_recommendation"]
            icons = {"Strong Hire": "🟢", "Hire": "🔵", "Borderline": "🟡", "No Hire": "🔴"}
            tr_map = {"Strong Hire": "Kesinlikle İşe Al", "Hire": "İşe Al", "Borderline": "Sınırda", "No Hire": "İşe Alma"}
            st.markdown(f"**{icons.get(rec, '⚪')} {tr_map.get(rec, rec)}**")

        st.markdown("---")
        st.markdown("### ⚡ İşlemler")
        if st.button("🔄 Mülakatı Sıfırla", use_container_width=True):
            for key in list(st.session_state.keys()):
                if key not in ["api_key", "model"]:
                    del st.session_state[key]
            init_session_state()
            st.rerun()

        st.markdown("---")
        st.markdown(
            '<div style="text-align:center; color:#94a3b8; font-size:0.75rem;">AI Mülakat Simülatörü v1.0<br>OpenRouter ile Güçlendirilmiştir</div>',
            unsafe_allow_html=True
        )


def main() -> None:
    init_session_state()
    render_sidebar()

    page = st.session_state.get("page", "home")

    if page == "home":
        render_home_page()
    elif page == "interview":
        render_interview_page()
    elif page == "analysis":
        render_analysis_page()
    elif page == "report":
        render_report_page()
    else:
        st.session_state["page"] = "home"
        render_home_page()


if __name__ == "__main__":
    main()
